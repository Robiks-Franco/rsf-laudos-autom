"""
=================================================================
 Núcleo do Sistema de Automação de Laudos (lógica compartilhada)
 Ocular Oftalmologia - Governador Valadares/MG
=================================================================
 
Este módulo contém TODA a lógica de extração/validação/geração de
laudos, sem NENHUMA dependência de interface gráfica. Ele é usado por
dois "front-ends" diferentes:
 
    - main.py     -> aplicativo de mesa (janela Tkinter, roda no seu
                      computador Windows/Mac/Linux)
    - app_web.py   -> aplicativo web (FastAPI), acessível pelo
                       navegador em qualquer computador, celular ou
                       tablet, depois de publicado na internet
 
Isso evita duplicar a lógica de extração/geração em dois lugares —
qualquer melhoria feita aqui (ex: um novo campo, uma correção na
extração) vale automaticamente para os dois aplicativos.
 
Arquitetura (classes principais):
    - ExamConfig      : carrega a configuração do tipo de exame (JSON)
    - PDFExtractor     : lê o(s) PDF(s) e extrai os dados via Claude API
    - ValidadorDados    : valida os dados extraídos contra a configuração
    - WordGenerator      : preenche o template Word com os dados
    - GeradorLaudoOCT     : orquestra todo o fluxo (extrair -> validar -> gerar)
 
PRONTO PARA ADICIONAR NOVOS TIPOS DE EXAME:
    Para suportar um novo exame (ex: Campo Visual, Topografia, Biometria),
    basta criar um novo arquivo "config_<tipo>.json" seguindo a mesma
    estrutura de "config_oct.json" e um novo "template_<tipo>.docx".
    Nenhuma classe abaixo precisa ser alterada — todas recebem o
    caminho da configuração/template como parâmetro.
 
Instale as dependências com: pip install -r requirements.txt
=================================================================
"""
 
from __future__ import annotations
 
import os
import re
import json
import base64
from pathlib import Path
from datetime import datetime, date
 
# -----------------------------------------------------------------
# Importação das bibliotecas externas.
# Se alguma não estiver instalada, o erro só é lançado quando a
# funcionalidade correspondente for realmente usada — assim quem
# estiver chamando (janela Tkinter ou rota web) consegue mostrar uma
# mensagem amigável em vez de travar com um erro técnico do Python.
# -----------------------------------------------------------------
try:
    import fitz  # PyMuPDF -> leitura de PDF e conversão de páginas em imagem
except ImportError:
    fitz = None
 
try:
    import docx  # python-docx -> geração/edição do laudo Word
except ImportError:
    docx = None
 
try:
    import anthropic  # SDK oficial da Anthropic -> chamada à Claude API
except ImportError:
    anthropic = None
 
 
# ===================================================================
# EXCEÇÕES PERSONALIZADAS
# Facilitam mostrar mensagens de erro claras e específicas, em vez de
# um traceback genérico do Python (tanto na janela quanto na web).
# ===================================================================
class ErroLaudoOCT(Exception):
    """Classe base para todos os erros previstos deste sistema."""
    pass
 
 
class ErroConfiguracao(ErroLaudoOCT):
    """Erro ao carregar/ler o arquivo de configuração (config_*.json)."""
    pass
 
 
class ErroExtracaoPDF(ErroLaudoOCT):
    """Erro ao ler o(s) PDF(s) ou ao extrair dados via Claude API."""
    pass
 
 
class ErroValidacaoDados(ErroLaudoOCT):
    """Erro quando dados obrigatórios não foram encontrados/estão inválidos."""
    pass
 
 
class ErroGeracaoDocumento(ErroLaudoOCT):
    """Erro ao preencher ou salvar o documento Word final."""
    pass
 
 
# ===================================================================
# REGISTRO DE EQUIPAMENTOS SUPORTADOS
# Cada entrada aponta para o arquivo de configuração e o template Word
# daquele equipamento/exame. É usado pelos dois "front-ends" (main.py
# e app_web.py) para montar o seletor de equipamento — assim os dois
# aplicativos sempre mostram exatamente as mesmas opções.
#
# PRONTO PARA ADICIONAR UM NOVO EQUIPAMENTO (ex: Topcon Triton):
#   1. Crie "config_<equipamento>.json" e "template_laudo_<equipamento>.docx"
#      seguindo o padrão dos exemplos já existentes.
#   2. Adicione uma linha abaixo, com uma chave curta (ex: "topcon").
#   Pronto — o aplicativo de mesa e o aplicativo web já passam a
#   oferecer essa opção automaticamente, sem mais nenhuma alteração.
# ===================================================================
EQUIPAMENTOS_SUPORTADOS = {
    "zeiss": {
        "nome_exibicao": "Zeiss Cirrus HD-OCT",
        "config": "config_oct.json",
        "template": "template_laudo.docx",
        "banco_conclusoes": "banco_conclusoes_zeiss.docx",
    },
    "nidek": {
        "nome_exibicao": "Nidek RS-3000 Advance",
        "config": "config_nidek.json",
        "template": "template_laudo.docx",
        "banco_conclusoes": "banco_conclusoes_nidek.docx",
    },
    "topcon_maestro2_wide": {
        "nome_exibicao": "Topcon Maestro2 3D Wide Report",
        "config": "config_oct_topcon_maestro2_wide.json",
        "template": "template_laudo_topcon_maestro2_wide.docx",
        "banco_conclusoes": "banco_conclusoes_topcon_maestro2_wide.docx",
    },
    "pentacam": {
        "nome_exibicao": "Pentacam® (Oculus) — Tomografia de Córnea",
        "config": "config_pentacam.json",
        "template": "template_laudo_pentacam.docx",
        "banco_conclusoes": "banco_conclusoes_pentacam.docx",
    },
    "campo_visual": {
        "nome_exibicao": "Octopus 600 (Haag-Streit) — Campo Visual",
        "config": "config_campo_visual.json",
        "template": "template_laudo_campo_visual.docx",
        "banco_conclusoes": "banco_conclusoes_campo_visual.docx",
    },
}
# NOTA: o protocolo "topcon_maestro2_disc" foi removido deste projeto por
# decisão do usuário (foco só no 3D Wide Report, 2 PDFs por exame: OD + OE).
# Os arquivos config_oct_topcon_maestro2_disc.json e
# template_laudo_topcon_maestro2_disc.docx nunca chegaram a existir no
# repositório, então não há nada pra apagar além desta entrada do dict.
 
 
# ===================================================================
# FUNÇÕES UTILITÁRIAS (datas, idade, protocolo)
# Ficam fora das classes por serem funções puras/independentes,
# reaproveitáveis por qualquer parte do sistema (e fáceis de testar).
# ===================================================================
 
# Formatos de data aceitos na entrada. IMPORTANTE: datas como "03/04/1978"
# são AMBÍGUAS — podem ser 3 de abril (dd/mm) ou 4 de março (mm/dd) — e as
# duas leituras costumam gerar uma idade "plausível" (às vezes só 1 ano de
# diferença), então um erro de interpretação aqui não é óbvio de detectar
# depois. Por isso a ORDEM em que os formatos são tentados importa muito:
# cada equipamento tem um formato "nativo" conhecido (o que aparece
# impresso no próprio exame), guardado em "formato_data_origem" no
# config_*.json — ver ExamConfig.formato_data_origem. Tentamos esse
# formato primeiro; os demais ficam como reserva, caso a IA já tenha
# convertido a data por conta própria.
_FORMATOS_POR_ORIGEM = {
    # "MDY" = mês/dia/ano (padrão americano — ex: Zeiss Cirrus, Octopus 600)
    "MDY": ("%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"),
    # "DMY" = dia/mês/ano (padrão brasileiro — usado como padrão quando o
    # equipamento não declara "formato_data_origem" no config)
    "DMY": ("%d/%m/%Y", "%m/%d/%Y", "%Y-%m-%d", "%d-%m-%Y"),
}
 
 
def tentar_parsear_data(texto: str, formato_origem: str = "DMY"):
    """
    Tenta interpretar uma string de data, priorizando o formato nativo
    do equipamento de origem ('formato_origem': "MDY" ou "DMY") para
    resolver corretamente datas ambíguas como "03/04/1978". Retorna um
    objeto 'date' do Python em caso de sucesso, ou None se não conseguir.
    """
    if not texto or not isinstance(texto, str):
        return None
    texto = texto.strip()
    formatos = _FORMATOS_POR_ORIGEM.get(formato_origem, _FORMATOS_POR_ORIGEM["DMY"])
    for formato in formatos:
        try:
            return datetime.strptime(texto, formato).date()
        except ValueError:
            continue
    return None
 
 
def normalizar_data(texto: str, formato_origem: str = "DMY") -> str:
    """
    Converte uma data para o formato brasileiro dd/mm/aaaa, interpretando
    corretamente o formato nativo do equipamento de origem. Se não
    conseguir interpretar, devolve o texto original sem alteração (para
    não perder a informação).
    """
    data = tentar_parsear_data(texto, formato_origem)
    if data is None:
        return texto
    return data.strftime("%d/%m/%Y")
 
 
def calcular_idade(data_nascimento_texto: str, data_referencia_texto: str) -> str:
    """
    Calcula a idade (em anos completos) a partir da data de nascimento
    e da data do exame. Espera que ambas as datas já tenham sido
    normalizadas para dd/mm/aaaa (ver normalizar_data) antes de chegar
    aqui. Retorna uma string como "73 anos", ou None se não for possível
    calcular (datas ausentes/ilegíveis, ou idade fora de uma faixa
    plausível) — nesse caso o campo simplesmente fica em branco no
    laudo, sem travar o programa.
    """
    nascimento = tentar_parsear_data(data_nascimento_texto, "DMY")
    referencia = tentar_parsear_data(data_referencia_texto, "DMY") or date.today()
    if nascimento is None:
        return None
    anos = referencia.year - nascimento.year
    fez_aniversario = (referencia.month, referencia.day) >= (nascimento.month, nascimento.day)
    if not fez_aniversario:
        anos -= 1
    if anos < 0 or anos > 130:
        return None  # data provavelmente ilegível/errada — não arrisca mostrar um valor absurdo
    return f"{anos} anos"
 
 
# Lista padrão de palavras-chave, usada apenas como último recurso se
# um arquivo config_*.json não definir a própria lista em
# "protocolos_por_palavra_chave" (ver ExamConfig.protocolos_por_palavra_chave
# mais abaixo). Cada equipamento nomeia os arquivos de um jeito — por
# isso o ideal é sempre definir essa lista dentro do config de cada
# equipamento, não aqui no núcleo.
_PROTOCOLOS_PADRAO_FALLBACK = [
    ("Macular Thickness", "Espessura Macular (Retina)"),
    ("Ganglion Cell", "Complexo de Células Ganglionares (GCC)"),
    ("ONH and RNFL", "Nervo Óptico e RNFL (Glaucoma)"),
    ("HD 5 Line Raster", "Cortes de Alta Definição (HD Raster)"),
]
 
 
def detectar_protocolo(nomes_arquivos: list, mapa_protocolos: list = None) -> str:
    """
    Monta uma descrição textual de quais protocolos de varredura foram
    incluídos no exame, a partir dos nomes dos arquivos PDF selecionados
    (cada equipamento nomeia os arquivos de forma padronizada, incluindo
    o nome do protocolo). 'mapa_protocolos' é a lista de (palavra_chave,
    descrição) vinda do config_*.json do equipamento em uso — se não for
    informada, usa uma lista padrão de fallback. Se nada for reconhecido,
    retorna None.
    """
    mapa = mapa_protocolos or _PROTOCOLOS_PADRAO_FALLBACK
    protocolos_encontrados = []
    for nome in nomes_arquivos:
        for palavra_chave, descricao in mapa:
            if palavra_chave.lower() in nome.lower() and descricao not in protocolos_encontrados:
                protocolos_encontrados.append(descricao)
    if not protocolos_encontrados:
        return None
    return " + ".join(protocolos_encontrados)
 
 
def gerar_nome_base_arquivo(dados: dict) -> str:
    """
    Gera um nome de arquivo seguro (sem caracteres especiais) a partir
    do nome do paciente e da data do exame, usado para nomear o .docx
    e o .json de saída (já que não há mais um único PDF de origem).
    """
    nome_paciente = dados.get("nome_paciente") or "paciente"
    data_exame = dados.get("data_exame") or ""
 
    texto = f"{nome_paciente}_{data_exame}"
    texto = texto.strip().replace(" ", "_")
    # Remove qualquer caractere que não seja letra, número, underscore ou hífen
    texto = re.sub(r"[^A-Za-z0-9_\-]", "", texto)
    texto = re.sub(r"_+", "_", texto).strip("_")
    return texto or "laudo_oct"
 
 
def completar_campos_automaticos(dados: dict, nomes_arquivos: list, config: "ExamConfig") -> dict:
    """
    Preenche os campos que o programa calcula sozinho (não pedidos à
    IA): datas normalizadas, idade, protocolo do exame e dados do
    médico responsável (vindos do arquivo de configuração). Fica como
    função de módulo (em vez de método de GeradorLaudoOCT) para poder
    ser reaproveitada também por scripts de teste, sem precisar de
    chave de API nem de PDFs reais — ver exemplo/teste_geracao_laudo.py.
    """
    dados = dict(dados)
 
    formato_origem = config.formato_data_origem
    nascimento_bruto = dados.get("data_nascimento")
    exame_bruto = dados.get("data_exame")
 
    if nascimento_bruto:
        dados["data_nascimento"] = normalizar_data(nascimento_bruto, formato_origem)
    if exame_bruto:
        dados["data_exame"] = normalizar_data(exame_bruto, formato_origem)
 
    idade_calculada = calcular_idade(dados.get("data_nascimento"), dados.get("data_exame"))
 
    # Rede de segurança: se a idade deu implausível (None: negativa, maior
    # que 130 anos, ou data ilegível) usando o formato nativo esperado do
    # equipamento, tenta novamente assumindo que a IA já converteu a data
    # sozinha para o formato oposto (isso acontece às vezes, apesar da
    # instrução no prompt para transcrever a data literalmente). Só troca
    # se essa segunda tentativa realmente resultar em uma idade válida —
    # caso contrário mantém o resultado original (incluindo o "None").
    if idade_calculada is None and (nascimento_bruto or exame_bruto):
        formato_alternativo = "MDY" if formato_origem == "DMY" else "DMY"
        nascimento_alt = (
            normalizar_data(nascimento_bruto, formato_alternativo)
            if nascimento_bruto else dados.get("data_nascimento")
        )
        exame_alt = (
            normalizar_data(exame_bruto, formato_alternativo)
            if exame_bruto else dados.get("data_exame")
        )
        idade_alt = calcular_idade(nascimento_alt, exame_alt)
        if idade_alt is not None:
            dados["data_nascimento"] = nascimento_alt
            dados["data_exame"] = exame_alt
            idade_calculada = idade_alt
 
    if idade_calculada:
        dados["idade"] = idade_calculada
 
    protocolo = detectar_protocolo(nomes_arquivos, config.protocolos_por_palavra_chave)
    if protocolo:
        dados["protocolo"] = protocolo
 
    medico = config.medico_responsavel
    dados["nome_medico"] = medico.get("nome", "")
    dados["crm_medico"] = medico.get("crm", "")
    dados["data_geracao_laudo"] = datetime.now().strftime("%d/%m/%Y %H:%M")
 
    return dados
 
 
# ===================================================================
# CLASSE: ExamConfig
# Responsabilidade única: carregar e disponibilizar a configuração
# de um tipo de exame (campos, validações, prompt de extração).
# ===================================================================
class ExamConfig:
    """
    Carrega a configuração de um tipo de exame a partir de um arquivo
    JSON (ex: config_oct.json) e expõe métodos utilitários para as
    demais classes do sistema.
    """
 
    def __init__(self, caminho_config: str):
        self.caminho_config = Path(caminho_config)
        self.dados_config = self._carregar_config()
 
    def _carregar_config(self) -> dict:
        """Lê e valida o arquivo JSON de configuração."""
        if not self.caminho_config.exists():
            raise ErroConfiguracao(
                f"Arquivo de configuração não encontrado: '{self.caminho_config}'.\n"
                "Verifique se o arquivo config_oct.json está na mesma pasta do programa."
            )
        try:
            with open(self.caminho_config, "r", encoding="utf-8") as arquivo:
                config = json.load(arquivo)
        except json.JSONDecodeError as erro:
            raise ErroConfiguracao(
                f"O arquivo de configuração '{self.caminho_config.name}' possui um "
                f"JSON inválido (erro na linha {erro.lineno}). Verifique a formatação."
            )
        if "campos" not in config or not config["campos"]:
            raise ErroConfiguracao(
                f"O arquivo '{self.caminho_config.name}' não define nenhum campo em 'campos'."
            )
        return config
 
    # ---------------- Propriedades de acesso rápido ----------------
    @property
    def tipo_exame(self) -> str:
        return self.dados_config.get("tipo_exame", "DESCONHECIDO")
 
    @property
    def nome_exibicao(self) -> str:
        return self.dados_config.get("nome_exibicao", self.tipo_exame)
 
    @property
    def campos(self) -> list:
        return self.dados_config.get("campos", [])
 
    @property
    def modelo_claude(self) -> str:
        return self.dados_config.get("modelo_claude", "claude-sonnet-5")
 
    @property
    def opcoes_extracao(self) -> dict:
        return self.dados_config.get("opcoes_extracao", {})
 
    @property
    def medico_responsavel(self) -> dict:
        return self.dados_config.get("medico_responsavel", {"nome": "", "crm": ""})
 
    @property
    def protocolos_por_palavra_chave(self) -> list:
        """
        Lista de (palavra_chave, descrição) usada para reconhecer, pelo
        nome dos arquivos PDF, quais protocolos de varredura foram
        incluídos no exame — específica de cada equipamento/config.
        Formato no JSON: [["palavra", "descrição"], ...].
        """
        pares = self.dados_config.get("protocolos_por_palavra_chave", [])
        return [tuple(par) for par in pares]
 
    @property
    def palavras_chave_verificacao(self) -> list:
        """
        Lista de palavras-chave (em minúsculas) que devem aparecer no
        texto do campo 'equipamento' lido pela IA diretamente do exame,
        usada como confirmação de que os PDFs enviados realmente são
        deste equipamento — proteção contra o usuário selecionar o
        equipamento errado no seletor (ex: enviar PDFs do Pentacam com
        'Zeiss Cirrus' escolhido no aplicativo). Ver ValidadorDados.
        """
        return [p.lower() for p in self.dados_config.get("palavras_chave_verificacao", [])]
 
    @property
    def formato_data_origem(self) -> str:
        """
        Formato nativo em que ESTE equipamento imprime as datas no PDF do
        exame: "MDY" (mês/dia/ano — padrão americano, ex: Zeiss Cirrus e
        Octopus 600) ou "DMY" (dia/mês/ano — padrão brasileiro, usado como
        valor padrão quando o config não declara isso explicitamente).
 
        Isso existe porque datas como "03/04/1978" são ambíguas (podem
        ser 3 de abril ou 4 de março) e as duas leituras costumam gerar
        uma idade "plausível" — então, sem saber o formato nativo de cada
        equipamento, o sistema não tinha como saber qual interpretação
        estava correta, e já calculou idade errada por causa disso. Ver
        tentar_parsear_data/normalizar_data/completar_campos_automaticos.
        """
        valor = self.dados_config.get("formato_data_origem", "DMY")
        return valor if valor in ("MDY", "DMY") else "DMY"
 
    @property
    def quantidade_pdfs_esperada(self):
        """
        Se definido no config_*.json ("quantidade_pdfs_esperada": N), o
        sistema exige que o exame venha com exatamente N arquivos PDF —
        usado pelo Topcon Maestro2 Wide (sempre 2: Wide OD + Wide OE).
        Retorna None (sem restrição) se o config não definir isso, para
        não afetar equipamentos que aceitam um número variável de PDFs
        (ex: Zeiss Cirrus, que pode ter vários protocolos por exame).
        """
        valor = self.dados_config.get("quantidade_pdfs_esperada")
        return int(valor) if valor is not None else None
 
    @property
    def recortes_por_categoria(self) -> dict:
        """
        Recortes adicionais (regiões específicas da página) enviados como
        UMA IMAGEM EXTRA, em alta resolução, só na chamada de uma
        categoria específica — usado junto com
        'dividir_extracao_por_categoria', quando uma categoria tem um
        gráfico pequeno/denso demais dentro da imagem da página inteira
        (ex: a pizza de 6 setores do GCL+ do Topcon Maestro2, que ocupa
        uma fração pequena da página e fica difícil de ler quando a
        imagem inteira é redimensionada pela API).

        Formato no JSON (dentro de "opcoes_extracao"):
            "recortes_por_categoria": {
                "<categoria>": {
                    "regiao_pdf": [x0, y0, x1, y1],
                    "dpi_imagem": 600
                }
            }
        'regiao_pdf' são coordenadas em PONTOS PDF (não pixels — 1 ponto
        = 1/72 polegada), medidas a partir do canto superior esquerdo da
        primeira página. 'dpi_imagem' é opcional; se omitido, usa o
        dobro do 'dpi_imagem' padrão do equipamento.

        Retorna {} se não definido — não afeta equipamentos que não
        usarem essa opção.
        """
        return self.opcoes_extracao.get("recortes_por_categoria", {})

    @property
    def dividir_extracao_por_categoria(self) -> bool:
        """
        Se True, em vez de pedir TODOS os campos numa única chamada à IA,
        o sistema divide a extração em várias chamadas menores, uma por
        categoria de campo (ex: retina_analysis, ganglion_cell,
        nervo_optico, disco_topografia) — cada uma reaproveitando as
        mesmas imagens já convertidas, só mudando os campos pedidos.
 
        Isso existe porque, em equipamentos com MUITOS campos (ex: Topcon
        Maestro2, ~80 campos), pedir tudo de uma vez numa única chamada
        sobrecarrega o modelo: ele passa a ficar cauteloso demais e
        retorna null em campos que, na verdade, consegue ler perfeitamente
        — confirmado num teste manual comparando um pedido simples de 5
        campos (acertou os números) contra o pedido completo de 82 campos
        na MESMA imagem (retornou tudo null). Dividir por categoria reduz
        a quantidade de campos por chamada e resolve isso, ao custo de
        várias chamadas à API em vez de uma (mais lento e mais caro por
        laudo — mas só para os equipamentos que ligarem esta opção).
 
        Desligado (False) por padrão, para não alterar o comportamento
        de equipamentos que já funcionam bem com uma chamada só (Zeiss,
        Nidek, Pentacam, Campo Visual). Ligue com
        "dividir_extracao_por_categoria": true dentro de "opcoes_extracao"
        no config_*.json do equipamento que precisar (mesmo bloco onde
        já ficam "dpi_imagem", "max_paginas_por_pdf" etc.).
        """
        return bool(self.opcoes_extracao.get("dividir_extracao_por_categoria", False))
 
    # ---------------- Métodos utilitários ----------------
    def obter_campo(self, campo_id: str) -> dict:
        """Retorna a definição de um campo específico pelo seu id."""
        for campo in self.campos:
            if campo["id"] == campo_id:
                return campo
        return None
 
    def campos_obrigatorios(self) -> list:
        """Retorna a lista de ids dos campos marcados como obrigatórios."""
        return [campo["id"] for campo in self.campos if campo.get("obrigatorio")]
 
    def categorias_dos_campos(self) -> list:
        """
        Lista as categorias distintas dos campos ('categoria' no JSON),
        na ordem em que aparecem no config. Usada quando
        'dividir_extracao_por_categoria' está ligado, para saber em
        quantos grupos separar a extração. Campos sem 'categoria'
        definida caem no grupo "geral".
        """
        vistas = []
        for campo in self.campos:
            categoria = campo.get("categoria", "geral")
            if categoria not in vistas:
                vistas.append(categoria)
        return vistas
 
    def campos_da_categoria(self, categoria: str) -> list:
        """Retorna apenas os campos que pertencem a uma categoria específica."""
        return [campo for campo in self.campos if campo.get("categoria", "geral") == categoria]
 
    def montar_prompt_extracao(self, campos: list = None) -> str:
        """
        Monta o prompt textual enviado à Claude API, descrevendo com
        precisão cada campo que deve ser extraído do exame, e de qual
        protocolo/arquivo cada um normalmente é originado.
 
        Por padrão descreve TODOS os campos do config ('self.campos').
        Se 'campos' for informado (lista de definições de campo, como as
        que estão em self.campos), monta o prompt só para esse subconjunto
        — usado por PDFExtractor quando 'dividir_extracao_por_categoria'
        está ligado, para pedir um grupo de campos por vez em vez de
        todos de uma vez.
        """
        campos = campos if campos is not None else self.campos
        base = self.dados_config.get("prompt_extracao_base", "")
 
        linhas_campos = []
        exemplo_json = {}
        for campo in campos:
            descricao = f'- "{campo["id"]}": {campo["rotulo"]}'
            if campo.get("unidade"):
                descricao += f' (unidade: {campo["unidade"]})'
            if campo.get("origem"):
                descricao += f' (protocolo/arquivo de origem esperado: {campo["origem"]})'
            if campo.get("valores_permitidos"):
                valores = ", ".join(campo["valores_permitidos"])
                descricao += f' (valores possíveis: {valores})'
            if campo.get("exemplos"):
                exemplos = ", ".join(campo["exemplos"])
                descricao += f' (exemplos: {exemplos})'
            linhas_campos.append(descricao)
            exemplo_json[campo["id"]] = None
 
        campos_texto = "\n".join(linhas_campos)
        exemplo_json_texto = json.dumps(exemplo_json, ensure_ascii=False, indent=2)
 
        instrucoes_finais = (
            "\n\nRegras importantes:\n"
            "1. Responda ESTRITAMENTE com um objeto JSON válido, sem nenhum texto, "
            "comentário ou marcação (como ```) antes ou depois.\n"
            "2. Use exatamente os identificadores de campo indicados acima como chaves do JSON, "
            "e APENAS esses identificadores (não crie campos extras, não aninhe objetos/dicionários "
            "dentro de um único campo).\n"
            "3. Se um dado não estiver visível, não fizer parte de nenhum dos arquivos fornecidos, "
            "ou pertencer a um protocolo cujo arquivo não foi enviado, use o valor null — nunca "
            "invente ou estime um valor.\n"
            "4. Cada arquivo fornecido abaixo está identificado por 'Arquivo N de M: <nome_do_arquivo>'. "
            "Use o nome do arquivo (e o texto extraído dele) para saber de qual protocolo os dados de "
            "cada campo devem ser lidos, conforme indicado entre parênteses em cada campo acima.\n"
            "5. Números devem ser retornados sem unidade (a unidade já é conhecida pelo sistema) e "
            "usando ponto como separador decimal (ex: 0.63, não 0,63).\n"
            "6. Para os campos de data, transcreva EXATAMENTE os números como aparecem "
            "impressos no exame, na mesma ordem (dia/mês/ano ou mês/dia/ano, o que for "
            "impresso) e separados por '/'. NÃO tente reordenar ou converter dia e mês — "
            "mesmo que o resultado pareça estranho, apenas copie os números exibidos. "
            "O sistema já sabe o formato nativo deste equipamento e faz essa conversão "
            "sozinho depois, de forma confiável; se a IA tentar converter, corre o risco "
            "de errar em datas ambíguas (ex: '03/04/1978' pode ser 3 de abril ou 4 de "
            "março) e isso já causou cálculo de idade errado no passado.\n"
            "7. Não faça diagnóstico nem interpretação clínica: apenas transcreva os dados exibidos.\n\n"
            f"Formato exato esperado (chaves, com valores de exemplo nulos):\n{exemplo_json_texto}"
        )
 
        return (
            f"{base}\n\n"
            f"Extraia os seguintes campos do exame de {self.nome_exibicao}:\n"
            f"{campos_texto}"
            f"{instrucoes_finais}"
        )
 
 
# ===================================================================
# CLASSE: PDFExtractor
# Responsabilidade única: ler o(s) PDF(s) de um mesmo exame e obter
# os dados estruturados usando a Claude API (texto + imagens).
# ===================================================================
class PDFExtractor:
    """
    Lê um ou mais PDFs de um mesmo exame (texto selecionável + imagens
    das páginas) e envia esse conteúdo à Claude API para extração
    estruturada dos dados definidos em ExamConfig — numa única chamada
    (comportamento padrão) ou em várias chamadas menores, uma por
    categoria de campo, se 'dividir_extracao_por_categoria' estiver
    ligado no config do equipamento (ver ExamConfig).
    """
 
    def __init__(self, config: ExamConfig, chave_api: str = None):
        if fitz is None:
            raise ErroExtracaoPDF(
                "A biblioteca PyMuPDF não está instalada.\n"
                "Instale com: pip install -r requirements.txt"
            )
        if anthropic is None:
            raise ErroExtracaoPDF(
                "A biblioteca 'anthropic' não está instalada.\n"
                "Instale com: pip install -r requirements.txt"
            )
 
        self.config = config
        opcoes = config.opcoes_extracao
        self.dpi_imagem = opcoes.get("dpi_imagem", 200)
        self.max_paginas_por_pdf = opcoes.get("max_paginas_por_pdf", 2)
        self.max_tokens_resposta = opcoes.get("max_tokens_resposta", 3072)
        # As imagens são enviadas em JPEG (não PNG): mapas/gráficos de OCT
        # são cheios de gradientes coloridos, que o PNG comprime muito mal
        # (arquivos de 4-6 MB por página não é incomum) — isso pode fazer o
        # pedido ultrapassar o limite de tamanho da API da Anthropic quando
        # o exame tem vários PDFs (erro HTTP 413 "request_too_large"). Em
        # JPEG, com qualidade alta, o mesmo conteúdo fica 5-7x menor sem
        # perda perceptível de legibilidade dos números.
        self.qualidade_jpeg = opcoes.get("qualidade_jpeg", 85)
 
        chave = chave_api or os.environ.get("ANTHROPIC_API_KEY")
        if not chave:
            raise ErroExtracaoPDF(
                "Chave da API da Anthropic não encontrada.\n"
                "Defina a variável de ambiente ANTHROPIC_API_KEY ou informe a "
                "chave quando solicitado."
            )
        self.cliente = anthropic.Anthropic(api_key=chave)
 
    def _extrair_texto(self, caminho_pdf: Path) -> str:
        """Extrai o texto selecionável de todas as páginas do PDF."""
        try:
            paginas_texto = []
            with fitz.open(caminho_pdf) as documento:
                for pagina in documento:
                    paginas_texto.append(pagina.get_text())
            return "\n".join(paginas_texto).strip()
        except Exception as erro:
            raise ErroExtracaoPDF(f"Falha ao ler o texto do PDF '{caminho_pdf.name}': {erro}")
 
    def _converter_paginas_em_imagens(self, caminho_pdf: Path) -> list:
        """
        Converte as páginas do PDF em imagens JPEG (base64), pois a
        maior parte dos valores de um exame de OCT aparece dentro de
        mapas/gráficos coloridos (imagem), não como texto selecionável.
        JPEG é usado em vez de PNG por comprimir muito melhor esse tipo
        de conteúdo (gradientes de cor), mantendo o pedido dentro do
        limite de tamanho da API mesmo com vários PDFs no mesmo exame.
        """
        imagens_base64 = []
        try:
            with fitz.open(caminho_pdf) as documento:
                zoom = self.dpi_imagem / 72
                matriz = fitz.Matrix(zoom, zoom)
                for indice, pagina in enumerate(documento):
                    if indice >= self.max_paginas_por_pdf:
                        break
                    pixmap = pagina.get_pixmap(matrix=matriz)
                    imagem_bytes = pixmap.tobytes("jpg", jpg_quality=self.qualidade_jpeg)
                    imagens_base64.append(base64.b64encode(imagem_bytes).decode("utf-8"))
            return imagens_base64
        except Exception as erro:
            raise ErroExtracaoPDF(
                f"Falha ao converter as páginas do PDF '{caminho_pdf.name}' em imagem: {erro}"
            )
 
    def _converter_regiao_em_imagem(self, caminho_pdf: Path, regiao_pdf: list, dpi: int) -> str:
        """
        Renderiza apenas uma região retangular da PRIMEIRA página do PDF
        (coordenadas em pontos PDF: [x0, y0, x1, y1]) numa resolução
        alta, e devolve o resultado já em base64 JPEG. Usado para dar à
        IA uma versão bem ampliada de um gráfico pequeno específico,
        além da imagem da página inteira — ver ExamConfig.recortes_por_categoria.
        """
        try:
            with fitz.open(caminho_pdf) as documento:
                pagina = documento[0]
                zoom = dpi / 72
                matriz = fitz.Matrix(zoom, zoom)
                clip = fitz.Rect(*regiao_pdf)
                pixmap = pagina.get_pixmap(matrix=matriz, clip=clip)
                imagem_bytes = pixmap.tobytes("jpg", jpg_quality=self.qualidade_jpeg)
                return base64.b64encode(imagem_bytes).decode("utf-8")
        except Exception as erro:
            raise ErroExtracaoPDF(
                f"Falha ao recortar região da página do PDF '{caminho_pdf.name}': {erro}"
            )

    def _montar_conteudo_arquivos(self, caminhos: list) -> list:
        """
        Lê e converte todos os PDFs (texto + imagens) em blocos de
        conteúdo prontos para a mensagem da API. Fica separado do resto
        de 'extrair_dados' para poder ser montado UMA ÚNICA VEZ e
        reaproveitado em várias chamadas à API (uma por categoria de
        campo, quando 'dividir_extracao_por_categoria' está ligado) —
        sem reconverter os mesmos PDFs em imagem repetidamente.
        """
        conteudo = []
        total_arquivos = len(caminhos)
        for indice, caminho in enumerate(caminhos, start=1):
            texto_pdf = self._extrair_texto(caminho)
            imagens_base64 = self._converter_paginas_em_imagens(caminho)
 
            cabecalho_arquivo = f"\n--- Arquivo {indice} de {total_arquivos}: {caminho.name} ---"
            if texto_pdf:
                cabecalho_arquivo += f"\nTexto extraído (pode estar incompleto):\n{texto_pdf}"
            conteudo.append({"type": "text", "text": cabecalho_arquivo})
 
            for imagem_b64 in imagens_base64:
                conteudo.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/jpeg",
                        "data": imagem_b64,
                    },
                })
        return conteudo
 
    def _verificar_tamanho(self, conteudo_mensagem: list):
        """
        Verificação proativa de tamanho: a API da Anthropic rejeita
        pedidos acima de um certo tamanho (erro HTTP 413
        "request_too_large"). Em vez de deixar o usuário receber esse
        erro técnico sem explicação, avisamos antes com uma sugestão
        prática (reduzir a resolução das imagens ou dividir o exame em
        duas gerações). O limite abaixo (25 MB de conteúdo base64) fica
        com folga do limite real da API, já que o texto do prompt e a
        própria requisição HTTP somam um pouco mais por cima.
        """
        tamanho_estimado_bytes = sum(
            len(bloco.get("source", {}).get("data", "")) for bloco in conteudo_mensagem
            if bloco.get("type") == "image"
        )
        if tamanho_estimado_bytes > 25_000_000:
            raise ErroExtracaoPDF(
                f"Os PDFs selecionados, convertidos em imagem, somam cerca de "
                f"{tamanho_estimado_bytes / 1_000_000:.0f} MB — isso provavelmente "
                "ultrapassa o limite de tamanho de pedido da API da Anthropic e "
                "resultaria em erro 'request_too_large'.\n\n"
                "Sugestões: (1) gere o laudo em duas etapas, selecionando menos "
                "PDFs de cada vez e completando os campos manualmente depois, ou "
                "(2) reduza o valor de 'dpi_imagem' no arquivo de configuração "
                "deste equipamento (ex: de 220 para 150) para diminuir o tamanho "
                "das imagens enviadas."
            )
 
    def _parsear_json(self, texto_resposta: str, motivo_parada: str = None) -> dict:
        """Extrai e valida o bloco JSON contido na resposta da IA."""
        texto = texto_resposta.strip()
 
        # Remove blocos de código markdown (```json ... ```), caso a IA os inclua
        if texto.startswith("```"):
            texto = texto.strip("`")
            if texto.lower().startswith("json"):
                texto = texto[4:]
 
        inicio = texto.find("{")
        fim = texto.rfind("}")
        if inicio == -1 or fim == -1:
            # Diagnóstico: mostra um trecho da resposta real da IA (ou avisa
            # se veio vazia), para dar pistas concretas do que aconteceu —
            # em vez de só uma mensagem genérica.
            if not texto:
                detalhe = "A IA devolveu uma resposta vazia (nenhum texto)."
            else:
                trecho = texto[:600]
                detalhe = f"Texto recebido da IA (trecho, para diagnóstico):\n\"{trecho}\""
 
            dica_parada = ""
            if motivo_parada and motivo_parada != "end_turn":
                dica_parada = (
                    f"\n\nMotivo de parada informado pela API: '{motivo_parada}'. "
                    + (
                        "Isso normalmente indica que a resposta foi cortada por atingir "
                        "o limite de tokens — tente novamente enviando menos PDFs de uma "
                        "vez, ou aumente 'max_tokens_resposta' no arquivo de configuração."
                        if motivo_parada == "max_tokens"
                        else "Verifique se o pedido não foi bloqueado por algum filtro de segurança."
                    )
                )
 
            raise ErroExtracaoPDF(
                "A resposta da IA não contém um JSON reconhecível. "
                "Tente novamente ou verifique a qualidade dos PDFs.\n\n"
                f"{detalhe}{dica_parada}"
            )
 
        bloco_json = texto[inicio: fim + 1]
        try:
            return json.loads(bloco_json)
        except json.JSONDecodeError as erro:
            trecho = bloco_json[-400:] if len(bloco_json) > 400 else bloco_json
            dica_parada = ""
            if motivo_parada == "max_tokens":
                dica_parada = (
                    "\n\nA resposta parece ter sido cortada por atingir o limite de "
                    "tokens (max_tokens) — tente novamente enviando menos PDFs de uma "
                    "vez, ou aumente 'max_tokens_resposta' no arquivo de configuração "
                    "deste equipamento."
                )
            raise ErroExtracaoPDF(
                f"Não foi possível interpretar o JSON retornado pela IA: {erro}\n\n"
                f"Final do texto recebido (para diagnóstico):\n\"...{trecho}\""
                f"{dica_parada}"
            )
 
    def _chamar_api(self, prompt: str, conteudo_arquivos: list) -> dict:
        """
        Monta a mensagem (texto do prompt + blocos de arquivo já
        convertidos) e faz UMA chamada à Claude API, devolvendo o JSON
        já interpretado. Separado de 'extrair_dados' para poder ser
        chamado várias vezes reaproveitando os mesmos 'conteudo_arquivos'
        (uma vez por categoria de campo, quando aplicável).
        """
        conteudo_mensagem = [{"type": "text", "text": prompt}] + conteudo_arquivos
        self._verificar_tamanho(conteudo_mensagem)
 
        try:
            resposta = self.cliente.messages.create(
                model=self.config.modelo_claude,
                max_tokens=self.max_tokens_resposta,
                messages=[{"role": "user", "content": conteudo_mensagem}],
            )
        except Exception as erro:
            raise ErroExtracaoPDF(
                f"Erro ao chamar a API da Claude: {erro}\n"
                "Verifique sua conexão com a internet e se a chave de API é válida."
            )
 
        texto_resposta = "".join(
            bloco.text for bloco in resposta.content if hasattr(bloco, "text")
        )
        motivo_parada = getattr(resposta, "stop_reason", None)
        return self._parsear_json(texto_resposta, motivo_parada=motivo_parada)
 
    def extrair_dados(self, caminhos_pdfs, callback_status=None) -> dict:
        """
        Fluxo completo de extração: lê um ou mais PDFs do mesmo exame,
        converte em imagens uma única vez, e chama a Claude API para
        extrair os dados definidos em ExamConfig — numa única chamada
        (comportamento padrão, igual antes) ou em várias chamadas
        menores, uma por categoria de campo, quando o config do
        equipamento tiver "dividir_extracao_por_categoria": true (ver
        ExamConfig.dividir_extracao_por_categoria para o motivo).
 
        'caminhos_pdfs' pode ser um único caminho (str) ou uma lista de
        caminhos — um exame de OCT normalmente é composto por vários PDFs.
        'callback_status', se informado, recebe mensagens de progresso
        (útil quando a extração é dividida em várias chamadas).
        """
        def status(mensagem):
            if callback_status:
                callback_status(mensagem)
 
        if isinstance(caminhos_pdfs, (str, Path)):
            caminhos_pdfs = [caminhos_pdfs]
        if not caminhos_pdfs:
            raise ErroExtracaoPDF("Nenhum arquivo PDF foi selecionado.")
 
        caminhos = [Path(caminho) for caminho in caminhos_pdfs]
        for caminho in caminhos:
            if not caminho.exists():
                raise ErroExtracaoPDF(f"Arquivo PDF não encontrado: '{caminho}'.")
            if caminho.suffix.lower() != ".pdf":
                raise ErroExtracaoPDF(f"O arquivo '{caminho.name}' não é um PDF.")
 
        conteudo_arquivos = self._montar_conteudo_arquivos(caminhos)
 
        if not self.config.dividir_extracao_por_categoria:
            prompt = self.config.montar_prompt_extracao()
            return self._chamar_api(prompt, conteudo_arquivos)
 
        # --- Extração dividida por categoria (uma chamada por grupo) ---
        categorias = self.config.categorias_dos_campos()
        recortes = self.config.recortes_por_categoria
        dados_combinados = {}
        for indice, categoria in enumerate(categorias, start=1):
            campos_categoria = self.config.campos_da_categoria(categoria)
            status(
                f"Extraindo grupo {indice} de {len(categorias)} ('{categoria}', "
                f"{len(campos_categoria)} campo(s))..."
            )
            prompt = self.config.montar_prompt_extracao(campos_categoria)

            # Começa com as imagens de página inteira (mesmas para todas as
            # categorias) e, se esta categoria tiver um recorte definido no
            # config, ACRESCENTA uma versão ampliada só daquela região —
            # sem alterar 'conteudo_arquivos' original, que é reaproveitado
            # pelas demais categorias.
            conteudo_desta_chamada = list(conteudo_arquivos)
            recorte = recortes.get(categoria)
            if recorte:
                dpi_recorte = recorte.get("dpi_imagem", self.dpi_imagem * 2)
                for caminho in caminhos:
                    imagem_recorte_b64 = self._converter_regiao_em_imagem(
                        caminho, recorte["regiao_pdf"], dpi_recorte
                    )
                    conteudo_desta_chamada.append({
                        "type": "text",
                        "text": (
                            f"\n--- Recorte ampliado de '{caminho.name}' (região específica "
                            "desta categoria, em resolução mais alta, para facilitar a "
                            "leitura de números pequenos) ---"
                        ),
                    })
                    conteudo_desta_chamada.append({
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": imagem_recorte_b64,
                        },
                    })

            resultado_categoria = self._chamar_api(prompt, conteudo_desta_chamada)
            dados_combinados.update(resultado_categoria)
        return dados_combinados
 
 
# ===================================================================
# CLASSE: ValidadorDados
# Responsabilidade única: validar os dados extraídos contra as
# regras definidas em ExamConfig (obrigatoriedade, valores permitidos).
# ===================================================================
class ValidadorDados:
    """Valida os dados extraídos de um exame contra a configuração."""
 
    def __init__(self, config: ExamConfig):
        self.config = config
 
    def validar(self, dados: dict) -> list:
        """
        Verifica os dados extraídos. Lança ErroValidacaoDados se algum
        campo obrigatório estiver ausente. Retorna uma lista de avisos
        (não bloqueantes) para inconsistências menores.
        """
        self._verificar_equipamento_correto(dados)
 
        campos_faltantes = []
        for campo_id in self.config.campos_obrigatorios():
            valor = dados.get(campo_id)
            if valor in (None, "", "null"):
                campo = self.config.obter_campo(campo_id)
                rotulo = campo["rotulo"] if campo else campo_id
                campos_faltantes.append(rotulo)
 
        if campos_faltantes:
            lista = ", ".join(campos_faltantes)
            raise ErroValidacaoDados(
                f"Não foi possível identificar os seguintes campos obrigatórios no "
                f"exame: {lista}. Verifique a qualidade dos PDFs ou preencha manualmente "
                "no laudo gerado."
            )
 
        avisos = []
        for campo in self.config.campos:
            valor = dados.get(campo["id"])
            valores_permitidos = campo.get("valores_permitidos")
            if valor and valores_permitidos:
                if str(valor).strip().upper() not in [v.upper() for v in valores_permitidos]:
                    avisos.append(
                        f"Campo '{campo['rotulo']}' retornou o valor '{valor}', fora do "
                        f"esperado {valores_permitidos}. Confira manualmente."
                    )
        return avisos
 
    def _verificar_equipamento_correto(self, dados: dict):
        """
        Confere se o texto de identificação do equipamento, lido pela IA
        diretamente das imagens do exame (campo 'equipamento'), é
        compatível com o equipamento selecionado no aplicativo. Isso
        evita que um exame de um equipamento (ex: Pentacam) seja
        processado com o modelo/config de outro (ex: Zeiss Cirrus) só
        porque o usuário esqueceu de trocar o seletor — nesse caso, os
        campos clínicos ficam todos em branco no laudo, o que é confuso
        e só é percebido depois de já ter sido gerado.
 
        Se o exame não tiver essa informação legível (campo vazio/null),
        a verificação é pulada — não bloqueia o fluxo normal por causa
        de uma imagem de baixa qualidade.
        """
        palavras_chave = self.config.palavras_chave_verificacao
        if not palavras_chave:
            return
 
        texto_equipamento = str(dados.get("equipamento") or "").strip()
        if not texto_equipamento or texto_equipamento.lower() == "null":
            return
 
        texto_lower = texto_equipamento.lower()
        if any(palavra in texto_lower for palavra in palavras_chave):
            return
 
        raise ErroValidacaoDados(
            f"Os PDFs selecionados podem não ser do equipamento escolhido "
            f"('{self.config.nome_exibicao}'). O texto de identificação do "
            f"equipamento lido dentro do próprio exame foi: '{texto_equipamento}'.\n\n"
            "Verifique se você selecionou o equipamento correto no seletor do "
            "aplicativo antes de gerar o laudo (ex: se o exame é do Pentacam, "
            "selecione 'Pentacam® (Oculus)', não 'Zeiss Cirrus HD-OCT' ou "
            "'Nidek RS-3000 Advance'), e gere o laudo novamente."
        )
 
 
# ===================================================================
# CLASSE: WordGenerator
# Responsabilidade única: preencher o template Word com os dados
# já extraídos e validados, gerando o laudo final.
# ===================================================================
class WordGenerator:
    """
    Preenche um template .docx substituindo marcadores no formato
    {{id_do_campo}} pelos valores extraídos do exame.
    """
 
    def __init__(self, caminho_template: str):
        if docx is None:
            raise ErroGeracaoDocumento(
                "A biblioteca python-docx não está instalada.\n"
                "Instale com: pip install -r requirements.txt"
            )
        self.caminho_template = Path(caminho_template)
        if not self.caminho_template.exists():
            raise ErroGeracaoDocumento(
                f"Template Word não encontrado: '{self.caminho_template}'.\n"
                "Verifique se o arquivo template_laudo.docx está na mesma pasta do programa."
            )
 
    def _valor_para_texto(self, valor) -> str:
        """
        Converte um valor extraído em texto para exibição no laudo.
        Valores ausentes (None/vazio) aparecem como travessão "—", para
        não deixar células em branco de forma confusa no documento.
        """
        if valor is None:
            return "—"
        texto = str(valor).strip()
        return texto if texto and texto.lower() != "null" else "—"
 
    def _aplicar_substituicoes(self, texto: str, dados: dict) -> str:
        """Substitui todos os marcadores {{campo}} presentes em um texto."""
        resultado = texto
        for chave, valor in dados.items():
            marcador = "{{" + chave + "}}"
            if marcador in resultado:
                resultado = resultado.replace(marcador, self._valor_para_texto(valor))
 
        # Rede de segurança: se sobrar algum marcador {{...}} sem
        # correspondência em 'dados' (ex: a IA não retornou aquele
        # campo, ou o protocolo não foi reconhecido), troca por "—" em
        # vez de deixar o texto "{{campo}}" cru e confuso no laudo.
        if "{{" in resultado:
            resultado = re.sub(r"\{\{[a-zA-Z0-9_]+\}\}", "—", resultado)
        return resultado
 
    def _substituir_em_paragrafo(self, paragrafo, dados: dict):
        """
        Substitui marcadores dentro de um parágrafo, preservando a
        formatação do primeiro "run" (trecho de texto) do parágrafo.
        """
        texto_completo = "".join(run.text for run in paragrafo.runs)
        if "{{" not in texto_completo:
            return
        novo_texto = self._aplicar_substituicoes(texto_completo, dados)
        if novo_texto == texto_completo or not paragrafo.runs:
            return
        paragrafo.runs[0].text = novo_texto
        for run in paragrafo.runs[1:]:
            run.text = ""
 
    def gerar_laudo(self, dados: dict, caminho_saida: str):
        """Gera o documento Word final preenchido e salva no caminho indicado."""
        try:
            documento = docx.Document(self.caminho_template)
        except Exception as erro:
            raise ErroGeracaoDocumento(f"Falha ao abrir o template Word: {erro}")
 
        # Substitui marcadores em parágrafos "soltos" do documento
        for paragrafo in documento.paragraphs:
            self._substituir_em_paragrafo(paragrafo, dados)
 
        # Substitui marcadores dentro de tabelas (identificação, parâmetros etc.)
        for tabela in documento.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        self._substituir_em_paragrafo(paragrafo, dados)
 
        caminho_saida = Path(caminho_saida)
        caminho_saida.parent.mkdir(parents=True, exist_ok=True)
        try:
            documento.save(caminho_saida)
        except PermissionError:
            raise ErroGeracaoDocumento(
                f"Não foi possível salvar '{caminho_saida.name}'. Feche o arquivo caso "
                "ele já esteja aberto no Word e tente novamente."
            )
        except Exception as erro:
            raise ErroGeracaoDocumento(f"Falha ao salvar o documento Word: {erro}")
 
 
# ===================================================================
# CLASSE: GeradorLaudoOCT (Orquestrador)
# Junta ExamConfig + PDFExtractor + ValidadorDados + WordGenerator
# em um fluxo único, do(s) PDF(s) de entrada ao laudo Word de saída.
#
# PRONTO PARA ADICIONAR NOVOS TIPOS DE EXAME: esta mesma classe pode
# ser reaproveitada para outro exame apenas trocando os caminhos de
# configuração e template passados no construtor.
# ===================================================================
class GeradorLaudoOCT:
    """Orquestra o fluxo completo: extração -> validação -> geração do laudo."""
 
    def __init__(self, caminho_config: str, caminho_template: str, chave_api: str = None):
        self.config = ExamConfig(caminho_config)
        self.extrator = PDFExtractor(self.config, chave_api=chave_api)
        self.validador = ValidadorDados(self.config)
        self.gerador_word = WordGenerator(caminho_template)
 
    def processar_exame(self, caminhos_pdfs, pasta_saida: str, callback_status=None) -> dict:
        """
        Executa o fluxo completo para um exame (um ou mais PDFs do mesmo
        paciente/exame) e retorna um dicionário com os caminhos gerados,
        avisos e dados extraídos.
        """
        def status(mensagem):
            if callback_status:
                callback_status(mensagem)
 
        if isinstance(caminhos_pdfs, (str, Path)):
            caminhos_pdfs = [caminhos_pdfs]
        caminhos_pdfs = [Path(caminho) for caminho in caminhos_pdfs]
 
        quantidade_esperada = self.config.quantidade_pdfs_esperada
        if quantidade_esperada is not None and len(caminhos_pdfs) != quantidade_esperada:
            raise ErroValidacaoDados(
                f"O {self.config.nome_exibicao} espera exatamente {quantidade_esperada} "
                f"arquivo(s) PDF por exame, mas {len(caminhos_pdfs)} foram selecionados. "
                "Confira se você enviou os arquivos certos (ex: um PDF do olho direito "
                "e um do olho esquerdo, ambos do protocolo 3D Wide Report) e tente novamente."
            )
 
        status(
            f"Lendo {len(caminhos_pdfs)} arquivo(s) PDF e extraindo dados com a "
            "Claude API (pode levar alguns segundos)..."
        )
        dados = self.extrator.extrair_dados(caminhos_pdfs, callback_status=callback_status)
 
        status("Calculando idade, protocolo e demais campos automáticos...")
        nomes_arquivos = [caminho.name for caminho in caminhos_pdfs]
        dados = completar_campos_automaticos(dados, nomes_arquivos, self.config)
 
        status("Validando os dados extraídos...")
        avisos = self.validador.validar(dados)
 
        nome_base = gerar_nome_base_arquivo(dados)
        pasta_saida = Path(pasta_saida)
        pasta_saida.mkdir(parents=True, exist_ok=True)
 
        caminho_json = pasta_saida / f"{nome_base}_dados.json"
        caminho_docx = pasta_saida / f"{nome_base}_laudo.docx"
 
        status("Salvando dados extraídos em JSON (backup/auditoria)...")
        with open(caminho_json, "w", encoding="utf-8") as arquivo:
            json.dump(dados, arquivo, ensure_ascii=False, indent=2)
 
        status("Preenchendo o laudo em Word...")
        self.gerador_word.gerar_laudo(dados, caminho_docx)
 
        status("Laudo gerado com sucesso!")
        return {
            "caminho_docx": str(caminho_docx),
            "caminho_json": str(caminho_json),
            "avisos": avisos,
            "dados": dados,
            "arquivos_processados": nomes_arquivos,
        }